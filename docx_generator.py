from docx import Document
import os
import sys
import math


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


TEMPLATE_PATH = resource_path("assets/Progress Report.docx")


# ---------- SAFE HELPERS ----------

def is_nan(value):
    return value is None or (isinstance(value, float) and math.isnan(value))


def normalize_mark(value):
    """
    Returns:
    - 'AB' if Ab / ab / absent
    - '0' if numeric zero
    - numeric string if valid number
    - '' if empty
    """
    if is_nan(value):
        return ""

    value = str(value).strip()

    if value.lower() in ["ab", "absent"]:
        return "AB"

    try:
        return str(int(float(value)))
    except Exception:
        return ""


def calculate_grade(mark):
    if mark == "AB" or mark == "":
        return "AB"

    mark = int(mark)

    if mark >= 22:
        return "A+"
    elif mark >= 18:
        return "A"
    elif mark >= 15:
        return "B"
    elif mark >= 10:
        return "C"
    else:
        return "D"


def safe_int(value, default=0):
    if is_nan(value):
        return default
    try:
        return int(float(value))
    except Exception:
        return default


# ---------- PLACEHOLDER REPLACEMENT ----------

def replace_placeholders_in_paragraph(paragraph, replacements):
    for run in paragraph.runs:
        if run.element.xpath(".//w:drawing"):
            return

    full_text = "".join(run.text for run in paragraph.runs)

    for key, val in replacements.items():
        full_text = full_text.replace(key, str(val))

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_everywhere(doc, replacements):
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, replacements)


# ---------- MAIN GENERATOR ----------

def generate_student_docx(student, report_title, output_folder="output/DOCX"):
    os.makedirs(output_folder, exist_ok=True)

    doc = Document(TEMPLATE_PATH)

    # Normalize marks
    eng = normalize_mark(student.get("english"))
    math = normalize_mark(student.get("math"))
    sci = normalize_mark(student.get("science"))

    # Totals
    total_marks = sum(
        int(m) for m in [eng, math, sci] if m not in ["", "AB"]
    )

    percentage = round((total_marks / 75) * 100, 2) if total_marks > 0 else 0

    total_classes = safe_int(student.get("total_classes"))
    classes_attended = safe_int(student.get("classes_attended"))

    attendance_percentage = (
        round((classes_attended / total_classes) * 100, 2)
        if total_classes > 0 else 0
    )

    replacements = {
        "{{student_name}}": student.get("student_name", ""),
        "{{father_name}}": student.get("father_name", ""),
        "{{mother_name}}": student.get("mother_name", ""),
        "{{gender}}": student.get("gender", ""),
        "{{age}}": safe_int(student.get("age"), ""),
        "{{standard}}": student.get("standard", ""),
        "{{roll_no}}": student.get("roll_no", ""),
        "{{address}}": student.get("address", ""),
        "{{year}}": student.get("year", ""),
        "{{contact_number}}": student.get("contact_number", ""),
        "{{report_title}}": report_title,

        "{{total_marks}}": total_marks,
        "{{percentage}}": f"{percentage}%",
        "{{total_classes}}": total_classes,
        "{{classes_attended}}": classes_attended,
        "{{attendance_percentage}}": f"{attendance_percentage}%",
    }

    replace_everywhere(doc, replacements)

    # SUBJECT TABLE
    for table in doc.tables:
        for row in table.rows:
            subject = row.cells[0].text.strip().lower()

            if subject == "english":
                row.cells[1].text = "25"
                row.cells[2].text = eng
                row.cells[3].text = calculate_grade(eng)

            elif subject in ["maths", "mathematics", "math"]:
                row.cells[1].text = "25"
                row.cells[2].text = math
                row.cells[3].text = calculate_grade(math)

            elif subject == "science":
                row.cells[1].text = "25"
                row.cells[2].text = sci
                row.cells[3].text = calculate_grade(sci)

    filename = f"{student['student_name'].replace(' ', '_')}.docx"
    path = os.path.join(output_folder, filename)
    doc.save(path)

    return path
